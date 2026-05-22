"""Generate docs/Design_Build_Roadmap.docx from roadmap content."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "Design_Build_Roadmap.docx"


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_checklist(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"[ ] {item}")


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build() -> None:
    doc = Document()
    set_doc_defaults(doc)

    # Cover
    t = doc.add_heading("Design Build Roadmap", 0)
    t.alignment = 0
    doc.add_paragraph("Wireframe to Production")
    doc.add_paragraph()
    meta = [
        "Project: Jawad Jalal - Portfolio and Landing Page Specialist",
        "Brand: Cinematic Mission (dark-first, strict monochrome, editorial halftone)",
        "Status: Planning locked, pre-build implementation",
        "Last updated: May 2026",
        "Companion: docs/design-build-roadmap.md",
    ]
    for line in meta:
        doc.add_paragraph(line)
    page_break(doc)

    # Executive summary
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(
        "This roadmap defines a six-stage pipeline from low-fidelity wireframes "
        "through production launch. Each stage has tools, outputs, exit criteria, "
        "and mapping to the Next.js 14 stack in this repository."
    )
    add_table(
        doc,
        ["Stage", "Name", "Primary output", "Primary tools"],
        [
            ["0a", "Wireframe v1", "Structure frames (all routes)", "HTML wireframe specs in docs/"],
            ["0b", "Wireframe v2", "Interaction annotations", "HTML wireframes v2"],
            ["0c", "Wireframe v3", "Build-ready handoff", "Approved wireframes"],
            ["1", "Mid-fi", "Full site skeleton in code", "Next.js, Tailwind, shadcn"],
            ["2", "Hi-fi", "Experience tier", "Framer Motion, GSAP, Lenis, Spline"],
            ["3", "Production maturity", "Launch-ready site", "Supabase, Resend, Vercel"],
        ],
    )
    doc.add_paragraph(
        "Gate rule: Do not start the next stage until the current stage passes its exit checklist."
    )
    page_break(doc)

    # Naming glossary
    doc.add_heading("Naming Glossary (Read First)", 1)
    doc.add_paragraph(
        "Two different uses of v1 appear in this project. Do not conflate them."
    )
    add_table(
        doc,
        ["Term", "Meaning", "Where defined"],
        [
            ["Wireframe v1", "First low-fi pass: IA, section order, grayscale blocks", "This roadmap, Stage 0a"],
            ["Wireframe v2", "Second pass: interactions, states, field schemas", "This roadmap, Stage 0b"],
            ["Wireframe v3", "Third pass: spacing, type labels, build handoff", "This roadmap, Stage 0c"],
            [
                "Launch content scope",
                "What ships at first launch (Option B work, no testimonials, interim wordmark)",
                "docs/launch-strategy-v1.md",
            ],
            ["Phase 2 (hero)", "Hero Blender video loop assets", "brand_assets/hero/README.md"],
        ],
    )
    page_break(doc)

    # Pipeline
    doc.add_heading("End-to-End Pipeline", 1)
    doc.add_paragraph(
        "Wireframe v1 -> Wireframe v2 -> Wireframe v3 -> Mid-fi -> Hi-fi -> Production maturity"
    )
    add_table(
        doc,
        ["From", "To", "Blocker if skipped"],
        [
            ["Wireframe v3", "Mid-fi", "Wrong layout, rework in code"],
            ["Mid-fi", "Hi-fi", "Motion on broken structure"],
            ["Hi-fi", "Production", "Launch without craft baseline"],
        ],
    )
    page_break(doc)

    # Stage 0
    doc.add_heading("Stage 0 - Low-Fi Wireframes (v1 to v3)", 1)
    doc.add_paragraph(
        "Purpose: Lock layout, section order, copy placement, and responsive behavior "
        "before production UI. Breakpoints: 375px mobile, 1200px desktop."
    )

    doc.add_heading("Wireframe v1 - Structure Pass", 2)
    add_table(
        doc,
        ["Route", "Status", "v1 deliverable file"],
        [
            ["/ Homepage", "Done", "docs/homepage-wireframe-spec.html"],
            ["/work", "Not started", "docs/wireframes-v1/work-wireframe-v1.html"],
            ["/work/[slug]", "Not started", "docs/wireframes-v1/case-study-wireframe-v1.html"],
            ["/services", "Not started", "docs/wireframes-v1/services-wireframe-v1.html"],
            ["/about", "Not started", "docs/wireframes-v1/about-wireframe-v1.html"],
            ["/journal", "Not started", "docs/wireframes-v1/journal-index-wireframe-v1.html"],
            ["/journal/[slug]", "Not started", "docs/wireframes-v1/journal-article-wireframe-v1.html"],
            ["/contact", "Not started", "docs/wireframes-v1/contact-wireframe-v1.html"],
            ["/404", "Not started", "docs/wireframes-v1/404-wireframe-v1.html"],
        ],
    )
    doc.add_paragraph("Homepage locked order: nav -> hero -> trust -> work -> process -> tools -> cta -> footer")
    doc.add_heading("Wireframe v1 Exit Checklist", 3)
    add_checklist(
        doc,
        [
            "All routes have desktop (1200px) and mobile (375px) frames",
            "Homepage proof-first order unchanged",
            "Locked copy pasted (no rewrites)",
            "Global header/footer on every page frame",
            "Hero: spaceman right, content left",
        ],
    )

    doc.add_heading("Wireframe v2 - Interaction Pass", 2)
    doc.add_paragraph("Deliverables under docs/wireframes-v2/ for each route.")
    add_bullets(
        doc,
        [
            "Nav scroll: transparent over hero -> nav-blur on scroll",
            "Hero layers L0-L3 per homepage-hero-interaction-audit.md",
            "CTA hover/focus per cta-messaging-matrix.md",
            "Mobile drawer: CTA row first, then links",
            "Form fields per ia-sitemap Contact section",
        ],
    )
    doc.add_heading("Wireframe v2 Exit Checklist", 3)
    add_checklist(
        doc,
        [
            "Developer can implement mid-fi without open interaction questions",
            "Hero approved against brand_assets/hero/README.md",
            "All CTA labels match CTA matrix",
        ],
    )

    doc.add_heading("Wireframe v3 - Build-Ready Handoff", 2)
    doc.add_paragraph("Deliverables under docs/wireframes-v3/ (or v2 files with v3 changelog).")
    add_bullets(
        doc,
        [
            "Section IDs: #hero, #trust, #work, #process, #tools, #cta",
            "Type scale: Space Grotesk, Inter, JetBrains Mono kickers",
            "Deferred to hi-fi: Spline, hero video, GSAP scroll scenes",
            "Deferred to production: Calendly live, Supabase, testimonials",
        ],
    )
    doc.add_heading("Wireframe v3 Exit Checklist", 3)
    add_checklist(
        doc,
        [
            "All page-briefs wireframe handoff items checked",
            "No open layout questions",
            "Conversion review complete (homepage-conversion-flow.md)",
        ],
    )
    page_break(doc)

    # Stage 1 Mid-fi
    doc.add_heading("Stage 1 - Mid-Fi (Bulk of the Build)", 1)
    doc.add_paragraph(
        "Purpose: Real site in Next.js with monochrome tokens, all routes, locked copy, "
        "hero poster. GSAP, Lenis, Spline, and hero video excluded."
    )
    add_table(
        doc,
        ["Layer", "Technology", "Mid-fi scope"],
        [
            ["Framework", "Next.js 14 App Router", "Server Components default"],
            ["Styling", "Tailwind 3 + CSS vars", "Monochrome tokens in globals.css"],
            ["UI", "shadcn/ui + CVA", "Buttons, inputs, cards remapped"],
            ["Sections", "components/sections/", "One file per homepage section"],
            ["Motion", "Framer Motion minimal", "Section fade-in only"],
        ],
    )
    doc.add_heading("Implementation Slices", 2)
    add_bullets(
        doc,
        [
            "Slice 1: tailwind.config.ts, globals.css, layout fonts, lib/site.ts",
            "Slice 2: navbar (nav order, Book a Call), footer",
            "Slice 3: homepage sections + app/page.tsx",
            "Slice 4: inner pages (/work, /services, /about, /journal, /contact, 404)",
            "Slice 5: lib/content/*.ts, journal MDX stubs",
        ],
    )
    doc.add_heading("Mid-Fi Exit Checklist", 3)
    add_checklist(
        doc,
        [
            "All routes render at 375px and 1200px with locked copy",
            "Marketing UI monochrome (no violet primary)",
            "Screenshot compare vs wireframe v3",
            "npm run build succeeds",
        ],
    )
    page_break(doc)

    # Stage 2 Hi-fi
    doc.add_heading("Stage 2 - Hi-Fi (Experience Tier)", 1)
    add_table(
        doc,
        ["#", "Package", "Details"],
        [
            ["1", "Hero immersion", "Ken Burns, optional video, scrim, headline pull"],
            ["2", "Nav choreography", "transparent -> nav-blur, height shrink"],
            ["3", "Process timeline", "GSAP orbit connector"],
            ["4", "Work cards", "Halftone overlay, shadow hover"],
            ["5", "Tools grid", "Elevated cards, logo SVGs"],
            ["6", "Page motion", "Staggered section enter"],
            ["7", "Spline orbit", "Desktop only, lazy-load, off on mobile"],
        ],
    )
    doc.add_paragraph("Hard rules: animate transform and opacity only; never transition-all; honor prefers-reduced-motion.")
    doc.add_heading("Hi-Fi Exit Checklist", 3)
    add_checklist(
        doc,
        [
            "Matches design-token-showcase.html and prototypes/hero-halftone.html",
            "Hero contrast QA passed",
            "No CLS from motion or lazy 3D",
        ],
    )
    page_break(doc)

    # Stage 3 Production
    doc.add_heading("Stage 3 - Production Maturity", 1)
    add_table(
        doc,
        ["Area", "Tool", "Tasks"],
        [
            ["Contact", "Supabase + Server Actions", "Form persistence, Zod validation"],
            ["Email", "Resend", "app/api/contact/route.ts"],
            ["Booking", "Calendly", "/contact#book live embed"],
            ["SEO", "Metadata API", "sitemap.ts, robots.ts, seo-page-map"],
            ["Deploy", "Vercel", "Preview and production env vars"],
            ["Analytics", "GA4 or Plausible", "Hero CTA, book call, form submit"],
        ],
    )
    doc.add_heading("Production Exit Checklist", 3)
    add_checklist(
        doc,
        [
            "Calendly live",
            "Contact form delivers via Resend",
            "Custom monogram replaces interim wordmark",
            "Core Web Vitals green on homepage",
            "All pages in sitemap",
            "404 recovery tested",
        ],
    )
    page_break(doc)

    # Traceability
    doc.add_heading("Traceability - Homepage Sections", 1)
    add_table(
        doc,
        ["Section", "ID", "Mid-fi component", "Hi-fi", "Production"],
        [
            ["Navigation", "nav", "navbar.tsx", "Scroll blur", "Calendly"],
            ["Hero", "#hero", "sections/hero.tsx", "Video, Spline", "OG, LCP"],
            ["Trust", "#trust", "trust-strip.tsx", "-", "-"],
            ["Work", "#work", "selected-work.tsx", "Card motion", "Case studies"],
            ["Process", "#process", "process.tsx", "GSAP", "-"],
            ["Tools", "#tools", "tools-showcase.tsx", "Logo SVGs", "-"],
            ["Final CTA", "#cta", "final-cta.tsx", "-", "Analytics"],
            ["Footer", "-", "footer.tsx", "-", "Newsletter"],
        ],
    )

    doc.add_heading("Tech Stack Matrix", 1)
    add_table(
        doc,
        ["Tool", "Wireframe", "Mid-fi", "Hi-fi", "Production"],
        [
            ["HTML wireframes", "Primary", "Reference", "Reference", "Archive"],
            ["Next.js 14", "-", "Primary", "Primary", "Primary"],
            ["Tailwind", "-", "Primary", "Primary", "Primary"],
            ["Framer Motion", "-", "Light", "Primary", "Primary"],
            ["GSAP", "-", "-", "Primary", "Primary"],
            ["Supabase", "-", "-", "-", "Primary"],
            ["Vercel", "-", "Preview", "Preview", "Primary"],
        ],
    )

    doc.add_heading("Risk Register", 1)
    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ["Wireframe vs launch scope confusion", "Use naming glossary; wireframes-vN folders"],
            ["Mid-fi before v3 approval", "Gate on v3 checklist"],
            ["Spline/video on mobile", "Desktop-only 3D; poster fallback"],
            ["Violet scaffold tokens", "Foundation slice before sections"],
        ],
    )

    doc.add_heading("Suggested Timeline", 1)
    add_table(
        doc,
        ["Week", "Focus", "Stage"],
        [
            ["A", "Wireframe v1 all inner routes", "0a"],
            ["B", "Wireframe v2 all routes", "0b"],
            ["C", "Wireframe v3 sign-off", "0c"],
            ["D-E", "Mid-fi foundation + homepage", "1"],
            ["F", "Mid-fi inner pages + content", "1"],
            ["G-H", "Hi-fi packages + QA", "2"],
            ["I-J", "Production + deploy", "3"],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph("Full detail: docs/design-build-roadmap.md")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
